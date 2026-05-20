"""Convert a markdown file to an editable Word document.

Usage: python md_to_docx.py [path/to/file.md]
Defaults to final-report.md in this folder.
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).parent
SRC = Path(sys.argv[1]) if len(sys.argv) > 1 else HERE / "final-report.md"
if not SRC.is_absolute():
    SRC = HERE / SRC
DST = SRC.with_suffix(".docx")

INLINE_BOLD = re.compile(r"\*\*([^*]+)\*\*")
INLINE_CODE = re.compile(r"`([^`]+)`")


def add_runs(paragraph, text):
    """Parse **bold** and `code` inline markers and emit styled runs."""
    pos = 0
    # Combined pattern: bold OR code
    pattern = re.compile(r"\*\*([^*]+)\*\*|`([^`]+)`")
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            run = paragraph.add_run(m.group(1))
            run.bold = True
        else:
            run = paragraph.add_run(m.group(2))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            add_runs(cells[c_idx].paragraphs[0], val)
    doc.add_paragraph()


def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # Fenced code block
        if line.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            continue

        # Table
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|\s*$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip separator
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            add_table(doc, header, rows)
            continue

        # Bullet list
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            # nested bullets directly under
            while i < len(lines) and lines[i].startswith("   - "):
                sub = doc.add_paragraph(style="List Bullet 2")
                add_runs(sub, lines[i].strip()[2:])
                i += 1
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Regular paragraph (may span continuation lines until blank)
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
